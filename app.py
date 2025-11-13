from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime
from io import BytesIO
import os

app = Flask(__name__)
OUTPUT_FOLDER = "output"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Function to generate DOCX
def create_docx(data):
    doc = Document("cover_page.docx")

    def format_run(run, size_pt):
        run.font.size = Pt(size_pt)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.bold = False

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, val in data.items():
            placeholder = "{{" + key + "}}"
            if placeholder in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if placeholder in inline[i].text:
                        inline[i].text = inline[i].text.replace(placeholder, val.upper())
                        format_run(inline[i], 28 if key=="topic" else 20)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, val in data.items():
                        placeholder = "{{" + key + "}}"
                        if placeholder in paragraph.text:
                            inline = paragraph.runs
                            for i in range(len(inline)):
                                if placeholder in inline[i].text:
                                    inline[i].text = inline[i].text.replace(placeholder, val.upper())
                                    format_run(inline[i], 28 if key=="topic" else 20)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

@app.route("/", methods=["GET","POST"])
def index():
    if request.method == "POST":
        data = {k:v for k,v in request.form.items()}
        now = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{now}.docx"
        path = os.path.join(OUTPUT_FOLDER, filename)

        file_data = create_docx(data)
        with open(path, "wb") as f:
            f.write(file_data.read())
        file_data.seek(0)

        return send_file(file_data, as_attachment=True, download_name=filename)

    return render_template("form.html")

if __name__ == "__main__":
    app.run(debug=True)
